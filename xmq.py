'''
Created on 

@author: xumao
'''

import json
import requests
from docx import Document
from docx.shared import Inches
import datetime
from urllib import quote
import io
from requests.exceptions import ProxyError

def get_data(url):
    querystring = {"count":"20","scope":"digests"}
    
    headers = {
        'Accept': "application/json, text/plain, */*",
        'Accept-Encoding': "gzip, deflate, br",
        'Accept-Language': "zh-CN,zh;q=0.9,en;q=0.8",
        'Authorization': "FDA892C7-EA6C-C2AC-8E86-038BBBC74B3B",
        'Connection': "keep-alive",
        'Host': "api.zsxq.com",
        'Origin': "https://wx.zsxq.com",
        'Referer': "https://wx.zsxq.com/dweb/",
        'User-Agent': "Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36",
        'X-Request-Id': "4c101f2a-7912-9914-026c-c133f7b28e0e",
        'X-Version': "1.10.5",
        'Cache-Control': "no-cache"
        }
    
    proxy = {'http': 'http://10.144.1.10:8080', 'https': 'http://10.144.1.10:8080'}
    try:
        response = requests.request("GET", url, headers=headers, timeout=5,verify=False, proxies = proxy)
    except ProxyError:
        print "proxy error, retry %s"%url
        response = requests.request("GET", url, headers=headers, timeout=5,verify=False, proxies = proxy)        
    
    result = json.loads(response.text)
    print result
    topics = result['resp_data']['topics']
    for topic in topics:
        file.add_paragraph(topic['create_time'])
        type = topic["type"]
#         print topic[type]['text']
        try:
            if type == "q&a":
                type = "question"
                file.add_paragraph("question: %s"%topic['question']['text'])
                file.add_paragraph("answer: %s"%topic['answer']['text'])
            else:
                file.add_paragraph(topic[type]['text'])
        except Exception,e:
            print "exception occurs \n"
            print "type is %s \n"%type
            print topic
        if topic[type].has_key('images'):
            images = topic[type]["images"]
            for image in images:
                image_name = "%s.jpg"%image["image_id"]
                try:   
                    pic_response = requests.request("GET", image["large"]["url"], stream=True, proxies = proxy)
                    image_name = io.BytesIO(pic_response.content)
                    file.add_picture(image_name, width=Inches(5))
                except Exception, e:
                    print "exception occurs when handle image %s in %s: %s"%(image["large"]["url"],topic['create_time'],e)
        if topic.has_key('show_comments'):
            comments = topic['show_comments']
            file.add_paragraph("comments:")
            for comment in comments:
                if comment.has_key('repliee'):
                    comment_content = "%s replies to %s: %s" %(comment['owner']['name'],comment['repliee'] ['name'],comment['text'])
                else:
                    comment_content = "%s : %s"%(comment['owner']['name'],comment['text'])
                file.add_paragraph(comment_content)
                
    next_page = response.json().get('resp_data').get('topics')
    if next_page:
        create_time = next_page[-1].get('create_time')

        
        end_time = create_time[:20]+str(int(create_time[20:23])-1).zfill(3)+create_time[23:]
        
        if create_time[20:23] == '000':
            temp_time = datetime.datetime.strptime(create_time, "%Y-%m-%dT%H:%M:%S.%f+0800")
            temp_time += datetime.timedelta(seconds=-1)
            end_time = temp_time.strftime("%Y-%m-%dT%H:%M:%S") + '.999+0800'
        print end_time
        end_time = quote(end_time)
        print end_time
        next_url = start_url + '&end_time=' + end_time
        print(next_url)
        get_data(next_url)

if __name__=="__main__":
    url = ""
    get_data(url)
    
    
        
        
    
    
    

    
    
    

    
    